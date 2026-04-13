#!/usr/bin/env python3
"""
Generiert die Beispieldateien für den Windows 11 Refresher Workshop.
Ausgabe: docs/beispieldateien/kurs-dateien-modul1.zip
         docs/beispieldateien/kurs-dateien-modul3.zip
"""

import os
import zipfile
import requests
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT_DIR = os.path.join(BASE_DIR, "docs", "beispieldateien")
TMP_DIR = os.path.join(BASE_DIR, "tmp_build")
os.makedirs(OUT_DIR, exist_ok=True)
os.makedirs(TMP_DIR, exist_ok=True)


# ── Modul 1: Ferienfotos + Packliste ──────────────────────────────────────────

# Fotos von Unsplash (Unsplash License – kostenlose Nutzung für Workshop erlaubt)
FOTOS = [
    ("Strand_Sonnenuntergang.jpg",  "https://images.unsplash.com/photo-1507525428034-b723cf961d3e?w=800&q=60"),
    ("Strand_Wellen.jpg",           "https://images.unsplash.com/photo-1519046904884-53103b34b206?w=800&q=60"),
    ("Berge_Panorama.jpg",          "https://images.unsplash.com/photo-1464822759023-fed622ff2c3b?w=800&q=60"),
    ("Berge_Wanderweg.jpg",         "https://images.unsplash.com/photo-1501854140801-50d01698950b?w=800&q=60"),
    ("Altstadt_Gasse.jpg",          "https://images.unsplash.com/photo-1467269204594-9661b134dd2b?w=800&q=60"),
    ("Markt_Fruechte.jpg",          "https://images.unsplash.com/photo-1488459716781-31db52582fe9?w=800&q=60"),
]

def download_fotos():
    paths = []
    for filename, url in FOTOS:
        dest = os.path.join(TMP_DIR, filename)
        if not os.path.exists(dest):
            print(f"  Lade {filename} herunter …")
            r = requests.get(url, timeout=30)
            r.raise_for_status()
            with open(dest, "wb") as f:
                f.write(r.content)
        paths.append(dest)
    return paths


def create_packliste_docx():
    doc = Document()
    doc.add_heading("Packliste – Familienausflug", level=1)
    doc.add_paragraph("Erstellt von: (Ihr Name)", style="Normal")
    doc.add_paragraph("")
    doc.add_heading("Kleidung", level=2)
    for item in ["T-Shirts (3x)", "Pullover", "Regenjacke", "Bequeme Schuhe"]:
        p = doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("Verpflegung", level=2)
    for item in ["Wasserflasche", "Sandwiches", "Obst", "Snacks"]:
        p = doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("Sonstiges", level=2)
    for item in ["Sonnencreme", "Erste-Hilfe-Set", "Kamera", "Wanderkarte"]:
        p = doc.add_paragraph(item, style="List Bullet")
    path = os.path.join(TMP_DIR, "Packliste-Familie.docx")
    doc.save(path)
    print("  Packliste-Familie.docx erstellt")
    return path


def create_modul1_zip(foto_paths, packliste_path):
    zip_path = os.path.join(OUT_DIR, "kurs-dateien-modul1.zip")
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for p in foto_paths:
            zf.write(p, os.path.basename(p))
        zf.write(packliste_path, "Packliste-Familie.docx")
    print(f"  → {zip_path}")
    return zip_path


# ── Modul 3: Brief + PDF-Belege ───────────────────────────────────────────────

def create_brief_docx():
    doc = Document()
    doc.add_heading("Brief an den Treuhänder", level=1)
    doc.add_paragraph("")
    doc.add_paragraph("Muster, 13. April 2026", style="Normal")
    doc.add_paragraph("")
    doc.add_paragraph("Sehr geehrte Damen und Herren,", style="Normal")
    doc.add_paragraph("")
    doc.add_paragraph(
        "anbei sende ich Ihnen die gewünschten Unterlagen für die "
        "Steuererklärung 2025. Die Belege finden Sie als separate Dateien.",
        style="Normal"
    )
    doc.add_paragraph("")
    doc.add_paragraph("Mit freundlichen Grüssen", style="Normal")
    doc.add_paragraph("")
    doc.add_paragraph("(Ihr Name)", style="Normal")
    path = os.path.join(TMP_DIR, "Brief-an-Treuhänder.docx")
    doc.save(path)
    print("  Brief-an-Treuhänder.docx erstellt")
    return path


def create_beleg_pdf(filename, titel, betrag, datum, absender):
    path = os.path.join(TMP_DIR, filename)
    c = canvas.Canvas(path, pagesize=A4)
    width, height = A4
    c.setFont("Helvetica-Bold", 16)
    c.drawString(60, height - 80, titel)
    c.setFont("Helvetica", 11)
    c.drawString(60, height - 120, f"Datum: {datum}")
    c.drawString(60, height - 145, f"Absender: {absender}")
    c.drawString(60, height - 170, f"Betrag: CHF {betrag}")
    c.setFont("Helvetica", 9)
    c.drawString(60, 40, "Beispieldokument – Windows 11 Refresher Workshop – Pro Senectute")
    c.save()
    print(f"  {filename} erstellt")
    return path


def create_modul3_zip(brief_path, beleg1_path, beleg2_path):
    zip_path = os.path.join(OUT_DIR, "kurs-dateien-modul3.zip")
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.write(brief_path,  "Brief-an-Treuhänder.docx")
        zf.write(beleg1_path, "Beleg-Krankenkasse.pdf")
        zf.write(beleg2_path, "Beleg-Strom.pdf")
    print(f"  → {zip_path}")
    return zip_path


# ── Hauptprogramm ─────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("=== Modul 1 ===")
    fotos   = download_fotos()
    pkliste = create_packliste_docx()
    create_modul1_zip(fotos, pkliste)

    print("=== Modul 3 ===")
    brief  = create_brief_docx()
    beleg1 = create_beleg_pdf(
        "Beleg-Krankenkasse.pdf",
        "Krankenkasse Musterland – Jahresabrechnung 2025",
        "1'840.00", "31.12.2025", "Krankenkasse Musterland AG, 6000 Luzern"
    )
    beleg2 = create_beleg_pdf(
        "Beleg-Strom.pdf",
        "Stadtwerke Musterstadt – Stromrechnung 2025",
        "312.50", "15.01.2026", "Stadtwerke Musterstadt, 6001 Luzern"
    )
    create_modul3_zip(brief, beleg1, beleg2)

    print("\nFertig. Dateien in docs/beispieldateien/")
